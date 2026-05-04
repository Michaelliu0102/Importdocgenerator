"""
Word文档生成模块 — 申报要素

- 进口流程 `generate()`：从空白文档生成标题与正文，不读取 templates/ 下的申报要素模板文件。
- 出口流程 `generate_from_export_template()`：以 `export_templates/申报要素总汇.docx`（可在 YAML
  `export_declaration_element_templates` 中配置）为壳，保留首段样式后填入正文。
"""

from pathlib import Path
from typing import Dict, Any, List, Optional

from docx import Document

from item_declaration_mapper import (
    build_declaration_groups,
    use_item_mapping_enabled,
)
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _to_float(value):
    """将字符串/数字安全转换为 float，失败返回 None。"""
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    s = str(value).strip().replace(",", ".")
    if not s:
        return None
    try:
        return float(s)
    except (TypeError, ValueError):
        return None


def _calc_gram_weight(net_weight, quantity):
    """计算每平方米克重(g/m2): net_weight / (quantity * 1.42) * 1000"""
    nw = _to_float(net_weight)
    qty = _to_float(quantity)
    if nw is None or qty is None or qty == 0:
        return None
    result = nw / (qty * 1.42) * 1000
    if result <= 0:
        return None
    return round(result, 2)


class DeclarationElementsGenerator:
    """申报要素文档生成器"""

    def __init__(self, config_path: Optional[str] = None):
        base = Path(__file__).resolve().parent
        self.config_path = config_path or str(
            base / "data" / "supplier_product_mapping_import.yaml"
        )

    def generate(
        self,
        invoice_data: Dict[str, Any],
        product_info: Dict[str, Any],
        supplier_info: Dict[str, Any],
        output_path: str,
        product_info_map: Dict[str, Dict] = None,
    ):
        """
        进口用：新建 Word，不使用 export_templates/申报要素总汇.docx。

        product_info_map: optional {hide_type: product_info} for per-item HS/elements.
        """
        doc = Document()

        title = doc.add_heading("申报要素", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._append_declaration_elements_body(
            doc,
            invoice_data,
            product_info,
            supplier_info,
            product_info_map or {},
        )

        doc.save(output_path)
        return output_path

    def generate_from_export_template(
        self,
        template_path: str,
        invoice_data: Dict[str, Any],
        product_info: Dict[str, Any],
        supplier_info: Dict[str, Any],
        output_path: str,
        product_info_map: Dict[str, Dict] = None,
    ) -> str:
        """基于 export_templates/申报要素总汇.docx：保留首段标题与样式，替换正文为自动生成的申报要素。"""
        doc = Document(template_path)
        for p in list(doc.paragraphs)[1:]:
            el = p._element
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

        self._append_declaration_elements_body(
            doc,
            invoice_data,
            product_info,
            supplier_info,
            product_info_map or {},
        )

        doc.save(output_path)
        return output_path

    def _append_declaration_elements_body(
        self,
        doc: Document,
        invoice_data: Dict[str, Any],
        product_info: Dict[str, Any],
        supplier_info: Dict[str, Any],
        product_info_map: Dict[str, Dict],
    ):
        items = invoice_data.get("customs_items") or invoice_data.get("items", [])
        currency = invoice_data.get("currency", "EUR")
        invoice_no = invoice_data.get("invoice_no", "")
        net_weight = invoice_data.get("net_weight", "")

        doc.add_paragraph(f"发票号: {invoice_no}")
        doc.add_paragraph(f"供应商: {supplier_info.get('name', '')}")
        doc.add_paragraph("")

        if use_item_mapping_enabled(invoice_data):
            item_groups = build_declaration_groups(
                items,
                invoice_data,
                product_info,
                product_info_map or {},
                self.config_path,
            )
        else:
            item_groups = self._group_items_by_product(
                items, product_info, product_info_map or {}
            )

        for group_idx, group in enumerate(item_groups, 1):
            group_hs = group["hs_code"]
            group_name = group["product_name"]
            group_elements = group["declaration_elements"]
            group_items = group["items"]

            # 克重为空时自动按公式回填：net weight/(quantity*1.42)
            total_qty = 0.0
            for item in group_items:
                q = _to_float(item.get("quantity", ""))
                if q is not None:
                    total_qty += q
            gram_weight = _calc_gram_weight(net_weight, total_qty)

            filled_elements = dict(group_elements or {})
            for k, v in list(filled_elements.items()):
                key_has_gram = ("克重" in str(k)) or ("g/m2" in str(k).lower())
                if not key_has_gram:
                    continue
                val = "" if v is None else str(v).strip()
                if (not val or val == "克/平方米") and gram_weight is not None:
                    filled_elements[k] = f"{gram_weight} g/m2"

            doc.add_heading(f"商品 {group_idx}: {group_name}", level=2)
            doc.add_paragraph(f"HS Code: {group_hs}")
            doc.add_paragraph("")

            # 申报要素编号列表
            para = doc.add_paragraph("申报要素:")
            para.runs[0].bold = True

            for idx, (key, value) in enumerate(filled_elements.items(), 1):
                if value:
                    doc.add_paragraph(f"  {idx}: {key}: {value}")

            doc.add_paragraph("")

            doc.add_paragraph("")

    def _group_items_by_product(
        self,
        items: list,
        product_info: Dict[str, Any],
        product_info_map: Dict[str, Dict] = None,
    ) -> list:
        """按HS Code / hide_type分组商品"""
        if not product_info_map:
            product_info_map = {}

        if not items:
            hs = product_info.get("hs_code", "") if product_info else ""
            name = product_info.get("name", "") if product_info else ""
            decl = product_info.get("declaration_elements", {}) if product_info else {}
            return [{"hs_code": hs, "product_name": name,
                     "declaration_elements": decl, "items": []}]

        groups = {}
        for item in items:
            prefix = item.get("item_code_prefix", "")
            hide_type = item.get("hide_type", "")

            if prefix == "5012":
                key = "5603149000_汽车装饰用材料"
                pi = {"hs_code": "5603149000", "name": "汽车装饰用材料",
                      "declaration_elements": product_info.get("declaration_elements", {}) if product_info else {}}
            elif prefix in ("5015", "5030", "5010"):
                key = "5603149000_面料"
                pi = {"hs_code": "5603149000", "name": "面料",
                      "declaration_elements": product_info.get("declaration_elements", {}) if product_info else {}}
            elif hide_type and hide_type in product_info_map:
                cur_pi = product_info_map[hide_type]
                key = f"{cur_pi.get('hs_code', '')}_{hide_type}"
                pi = cur_pi
            else:
                pi = product_info or {}
                key = f"{pi.get('hs_code', '')}_{pi.get('name', '')}"

            if key not in groups:
                groups[key] = {
                    "hs_code": pi.get("hs_code", ""),
                    "product_name": pi.get("name", ""),
                    "declaration_elements": pi.get("declaration_elements", {}),
                    "items": []
                }
            groups[key]["items"].append(item)

        return list(groups.values())
