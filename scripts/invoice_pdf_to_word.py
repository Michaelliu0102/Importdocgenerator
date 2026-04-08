#!/usr/bin/env python3
"""
Convert a specific invoice PDF into an editable Word (.docx).

This script is tailored to the "CustInvc_*.pdf" layout:
- Extract seller/bill-to/ship-to addresses
- Extract invoice number, invoice date
- Extract table rows
- Extract subtotal/tax/total and due date (Total下面的金额和日期)
- Rebuild a clean docx without barcode and page footer (since we generate new content)
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pypdf import PdfReader


def _norm_spaces(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def _extract_first_match(pattern: str, text: str, flags: int = 0) -> Optional[str]:
    m = re.search(pattern, text, flags=flags)
    if not m:
        return None
    return m.group(1).strip()


def _parse_invoice_text(pdf_path: str) -> Dict:
    reader = PdfReader(pdf_path)
    pages_text = []
    for page in reader.pages:
        pages_text.append(page.extract_text() or "")
    text = "\n".join(pages_text)
    # pypdf sometimes injects a hidden glyph (e.g. "ĉ") into numeric lines.
    text = text.replace("ĉ", "")
    text = text.replace("\u00a0", " ")

    # Normalize line list
    # pypdf sometimes returns hidden glyphs like "ĉ" and may also glue tokens (e.g. "1MF")
    lines = [ln.replace("ĉ", "").strip() for ln in text.splitlines()]
    lines = [ln for ln in lines if ln]

    # Invoice number: "#32748"
    invoice_no = _extract_first_match(r"#\s*([0-9]+)", text)  # digits only
    invoice_no = f"#{invoice_no}" if invoice_no else ""

    # Invoice date: 26/3/2026
    invoice_date = _extract_first_match(r"\b(\d{1,2}/\d{1,2}/\d{4})\b", text)

    # Due date: "Due Date: 26/3/2026"
    due_date = _extract_first_match(r"Due Date:\s*(\d{1,2}/\d{1,2}/\d{4})", text)

    # Total amount: take the last "Total <num>" occurrence.
    total_amount = None
    for m in re.finditer(r"^\s*Total\s*([0-9][0-9,\.]*)\s*$", text, flags=re.MULTILINE):
        total_amount = m.group(1)
    if total_amount is None:
        # fallback: first "Total" with numeric
        total_amount = _extract_first_match(r"Total\s+([0-9][0-9,\.]*)", text)

    # Subtotal and Tax total (optional for layout)
    subtotal = _extract_first_match(r"^\s*Subtotal\s*([0-9][0-9,\.]*)\s*$", text, flags=re.MULTILINE)
    tax_total = _extract_first_match(r"^\s*Tax\s*Total\s*([0-9][0-9,\.]*)\s*$", text, flags=re.MULTILINE)
    if tax_total is None:
        tax_total = _extract_first_match(r"^\s*Tax\s*Total\s*([0-9][0-9,\.]*)", text, flags=re.MULTILINE)

    # Seller address: first few lines before "Invoice"
    seller_lines: List[str] = []
    try:
        idx_invoice = next(i for i, ln in enumerate(lines) if ln.lower() == "invoice")
        seller_lines = [_norm_spaces(x) for x in lines[:idx_invoice] if _norm_spaces(x)]
    except StopIteration:
        seller_lines = [_norm_spaces(x) for x in lines[:8] if _norm_spaces(x)]

    # Bill to / Ship to
    bill_to_lines: List[str] = []
    ship_to_lines: List[str] = []
    # Find "Bill To Ship To"
    bill_ship_idx = None
    for i, ln in enumerate(lines):
        if "Bill To" in ln and "Ship To" in ln:
            bill_ship_idx = i
            break
    if bill_ship_idx is not None:
        # Take lines until "Due Date:"
        end_idx = None
        for j in range(bill_ship_idx + 1, len(lines)):
            if lines[j].startswith("Due Date:"):
                end_idx = j
                break
        seg = lines[bill_ship_idx + 1 : end_idx] if end_idx is not None else lines[bill_ship_idx + 1 :]

        # Split by occurrences containing "Japan"
        japan_idxs = [k for k, s in enumerate(seg) if "Japan" in s]
        if japan_idxs:
            # First Japan ends Bill To
            first_end = japan_idxs[0]
            bill_seg = seg[: first_end + 1]
            bill_to_lines = [ln for ln in bill_seg if _norm_spaces(ln)]

            # Second Japan ends Ship To if exists
            if len(japan_idxs) >= 2:
                second_end = japan_idxs[1]
                ship_seg = seg[first_end + 1 : second_end + 1]
                # remove trailing total in "Japan 377,250"
                ship_to_lines = []
                for ln in ship_seg:
                    ln2 = ln
                    if "Japan" in ln2:
                        # Keep only the word "Japan" if the line contains extra numbers/suffix.
                        ln2 = "Japan"
                    ln2 = _norm_spaces(ln2)
                    if ln2:
                        ship_to_lines.append(ln2)
            else:
                # If only one block found, treat the next lines as ship-to (best effort)
                ship_to_lines = []

    # Table rows: find header line containing "# Item" and "Unit Price" and "Amount"
    header_idx = None
    for i, ln in enumerate(lines):
        if ln.startswith("#") and "Item" in ln and "Unit Price" in ln and "Amount" in ln:
            header_idx = i
            break
        if "Unit Price" in ln and "Tax Rate" in ln and "Amount" in ln and "Item" in ln:
            header_idx = i
            break

    items: List[Dict] = []
    if header_idx is not None:
        for k in range(header_idx + 1, len(lines)):
            ln = lines[k]
            if ln.startswith("Subtotal"):
                break
            # item line must have tax rate like "0%"
            if not re.search(r"\b\d+%?\b", ln) and "%" not in ln:
                continue
            # Clean and normalize common PDF extraction artifacts.
            #  - "1MF ..."  -> "1 MF ..."
            #  - "ĉ1,440" -> "1,440" (we already removed "ĉ" globally above, but keep best-effort)
            clean_ln = ln.replace("ĉ", "")
            clean_ln = re.sub(r"^(\d+)([A-Za-z])", r"\1 \2", clean_ln)

            # parse based on tokenization
            tokens = clean_ln.split()
            if not tokens:
                continue

            # Ensure first token begins with a number
            m0 = re.match(r"^(\d+)", tokens[0])
            if not m0:
                continue
            # find token with '%' (tax rate)
            tax_i = None
            for idx, t in enumerate(tokens):
                if "%" in t:
                    tax_i = idx
                    break
            if tax_i is None or tax_i < 2 or tax_i + 1 >= len(tokens):
                continue
            item_no = m0.group(1)
            # Clean price/amount tokens to keep only digits + separators
            unit_price = re.sub(r"[^0-9,\.]", "", tokens[tax_i - 1])
            tax_rate = tokens[tax_i]
            amount = re.sub(r"[^0-9,\.]", "", tokens[tax_i + 1])

            before = tokens[1 : tax_i - 1]  # includes item name + quantity + units
            if not before:
                continue

            # quantity is first numeric-like token in 'before'
            qty_idx = None
            for idx, t in enumerate(before):
                if re.match(r"^[0-9][0-9,\.]*$", t) or re.match(r"^[0-9]+(\.[0-9]+)?$", t):
                    qty_idx = idx
                    break
            if qty_idx is None:
                # fallback: use the first token
                qty_idx = 0
            quantity = before[qty_idx]
            item_name_tokens = before[:qty_idx]
            units_tokens = before[qty_idx + 1 :]
            item_name = " ".join(item_name_tokens).strip()
            units = " ".join(units_tokens).strip()

            # clean commas spacing
            item_name = _norm_spaces(item_name)
            units = _norm_spaces(units)

            items.append(
                {
                    "line_no": item_no,
                    "item": item_name,
                    "quantity": quantity,
                    "units": units,
                    "unit_price": unit_price,
                    "tax_rate": tax_rate,
                    "amount": amount,
                }
            )

    return {
        "seller_lines": seller_lines,
        "bill_to_lines": bill_to_lines,
        "ship_to_lines": ship_to_lines,
        "invoice_no": invoice_no,
        "invoice_date": invoice_date or "",
        "due_date": due_date or "",
        "subtotal": subtotal or "",
        "tax_total": tax_total or "",
        "total_amount": total_amount or "",
        "items": items,
    }


def _add_kv_paragraph(doc: Document, left: str, right: str) -> None:
    p = doc.add_paragraph()
    run_left = p.add_run(left)
    run_left.bold = True
    p.add_run(" ")
    p.add_run(right)


def build_word_from_parsed(parsed: Dict, output_path: str) -> None:
    doc = Document()

    # Basic page margins for "invoice look"
    try:
        for section in doc.sections:
            section.top_margin = section.top_margin  # keep defaults
    except Exception:
        pass

    # Seller block (top-left)
    if parsed.get("seller_lines"):
        p = doc.add_paragraph()
        for i, ln in enumerate(parsed["seller_lines"]):
            if i > 0:
                p.add_run("\n")
            p.add_run(ln)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()

    # Bill To / Ship To and invoice meta
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"

    # Header cells
    t.cell(0, 0).text = "Bill To"
    t.cell(0, 1).text = "Ship To"
    t.cell(1, 0).text = "\n".join(parsed.get("bill_to_lines", [])) if parsed.get("bill_to_lines") else ""
    t.cell(1, 1).text = "\n".join(parsed.get("ship_to_lines", [])) if parsed.get("ship_to_lines") else ""

    doc.add_paragraph()

    # Invoice meta
    if parsed.get("invoice_no"):
        _add_kv_paragraph(doc, "Invoice #", parsed["invoice_no"])
    if parsed.get("invoice_date"):
        _add_kv_paragraph(doc, "Invoice Date", parsed["invoice_date"])

    doc.add_paragraph()

    # Items table
    items = parsed.get("items", [])
    headers = ["#", "Item", "Quantity", "Units", "Unit Price", "Tax Rate", "Amount"]
    item_table = doc.add_table(rows=1, cols=len(headers))
    item_table.style = "Table Grid"
    item_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = item_table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for it in items:
        row_cells = item_table.add_row().cells
        row_cells[0].text = str(it.get("line_no", ""))
        row_cells[1].text = str(it.get("item", ""))
        row_cells[2].text = str(it.get("quantity", ""))
        row_cells[3].text = str(it.get("units", ""))
        row_cells[4].text = str(it.get("unit_price", ""))
        row_cells[5].text = str(it.get("tax_rate", ""))
        row_cells[6].text = str(it.get("amount", ""))

    doc.add_paragraph()

    # Totals: Total below amount and date (Due Date)
    if parsed.get("subtotal"):
        _add_kv_paragraph(doc, "Subtotal", parsed["subtotal"])
    if parsed.get("tax_total"):
        _add_kv_paragraph(doc, "Tax Total", parsed["tax_total"])

    if parsed.get("total_amount"):
        _add_kv_paragraph(doc, "Total", parsed["total_amount"])
    if parsed.get("due_date"):
        _add_kv_paragraph(doc, "Due Date", parsed["due_date"])

    doc.save(output_path)


def main():
    ap = argparse.ArgumentParser(description="Convert CustInvc PDF to editable Word")
    ap.add_argument("--input", required=True, help="Input invoice PDF path")
    ap.add_argument("--output", default="", help="Output docx path (default: input stem + .docx)")
    args = ap.parse_args()

    in_path = Path(args.input).expanduser().resolve()
    if not in_path.exists():
        raise FileNotFoundError(in_path)

    out_path = Path(args.output).expanduser().resolve() if args.output else in_path.with_suffix(".docx")
    parsed = _parse_invoice_text(str(in_path))
    build_word_from_parsed(parsed, str(out_path))
    print(f"Generated editable Word: {out_path}")


if __name__ == "__main__":
    main()

